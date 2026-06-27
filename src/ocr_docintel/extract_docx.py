# DOCX를 문서 순서대로 제목/리스트/강조/표를 보존해 추출하는 파서 (PRD FR-2, System Prompt §3.2)
from __future__ import annotations
from pathlib import Path

import docx
from docx.document import Document as _DocxDocument
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table as _Table
from docx.text.paragraph import Paragraph

from .ir import Document, Element, Source


def _iter_blocks(parent):
    """본문 자식을 문서 순서대로 Paragraph/Table로 산출(둘의 순서 보존)."""
    body = parent.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield _Table(child, parent)


def _runs_to_md(para) -> str:
    """런의 강조 스타일(bold/italic/underline)을 마크다운으로 변환."""
    out = []
    for run in para.runs:
        t = run.text
        if not t:
            continue
        if run.bold:
            t = f"**{t}**"
        if run.italic:
            t = f"*{t}*"
        if run.underline:
            t = f"<u>{t}</u>"
        out.append(t)
    return "".join(out) or para.text


def _heading_level(style_name: str):
    """'Heading N' 스타일 → 제목 레벨. Title → 1."""
    if not style_name:
        return 0
    s = style_name.lower()
    if s == "title":
        return 1
    if s.startswith("heading"):
        try:
            return min(int(s.replace("heading", "").strip()), 4)
        except ValueError:
            return 0
    return 0


def _list_indent(para) -> int:
    """목록 단락의 들여쓰기 레벨(numbering ilvl 또는 들여쓰기 기반)."""
    try:
        ppr = para._p.pPr
        if ppr is not None and ppr.numPr is not None and ppr.numPr.ilvl is not None:
            return int(ppr.numPr.ilvl.val)
    except Exception:
        pass
    return 0


def extract(path: str | Path) -> Document:
    """DOCX → Document(IR)."""
    path = Path(path)
    d = docx.Document(str(path))
    elements: list[Element] = []
    order = 0
    title = path.stem

    for block in _iter_blocks(d):
        if isinstance(block, _Table):
            rows = [[("" if c.text is None else c.text.replace("\n", "<br>")) for c in row.cells]
                    for row in block.rows]
            elements.append(Element(type="table", order=order, level=0, content=rows,
                                    source=Source()))
            order += 1
            continue

        text = block.text.strip()
        if not text:
            continue
        style = block.style.name if block.style else ""
        level = _heading_level(style)
        if level > 0:
            if level == 1 and title == path.stem:
                title = text
            elements.append(Element(type="heading", order=order, level=level,
                                    content=text, source=Source()))
        elif style and "list" in style.lower():
            indent = "    " * _list_indent(block)
            elements.append(Element(type="list", order=order, level=_list_indent(block),
                                    content=f"{indent}- {_runs_to_md(block)}", source=Source()))
        else:
            elements.append(Element(type="paragraph", order=order, level=0,
                                    content=_runs_to_md(block), source=Source()))
        order += 1

    return Document(document_title=title, source_format="docx", elements=elements)
