# 테스트용 골든셋 PDF를 reportlab으로 생성하는 픽스처
import pytest
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.platypus import Table


@pytest.fixture
def single_column_pdf(tmp_path):
    """제목 2계층 + 본문 단락 PDF (G1)."""
    p = tmp_path / "single_column.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    w, h = A4
    c.setFont("Helvetica-Bold", 24)
    c.drawString(72, h - 90, "Annual Report")
    c.setFont("Helvetica-Bold", 16)
    c.drawString(72, h - 140, "Section One")
    c.setFont("Helvetica", 11)
    c.drawString(72, h - 170, "This is the first body paragraph of the report.")
    c.drawString(72, h - 190, "It continues with more plain body text here.")
    c.save()
    return p


@pytest.fixture
def simple_table_pdf(tmp_path):
    """단순 표(헤더 1행 + 데이터 2행) PDF (G2)."""
    p = tmp_path / "simple_table.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    w, h = A4
    c.setFont("Helvetica-Bold", 22)
    c.drawString(72, h - 90, "Table Doc")
    data = [["Item", "2024", "2025"], ["Sales", "120", "150"], ["Cost", "80", "90"]]
    t = Table(data, colWidths=[120, 80, 80], rowHeights=24)
    t.setStyle(__import__("reportlab.lib.colors", fromlist=["black"]) and _grid_style())
    t.wrapOn(c, w, h)
    t.drawOn(c, 72, h - 230)
    c.save()
    return p


@pytest.fixture
def two_column_pdf(tmp_path):
    """좌/우 2단 본문 PDF (FR-1 다단). 읽기 순서는 좌단 전체 → 우단 전체."""
    p = tmp_path / "two_column.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    w, h = A4
    c.setFont("Helvetica", 11)
    # 좌단 (x=72)
    c.drawString(72, h - 120, "LEFT one")
    c.drawString(72, h - 140, "LEFT two")
    # 우단 (x=340) — 더 위쪽에서 시작해도 좌단이 먼저 와야 한다.
    c.drawString(340, h - 100, "RIGHT one")
    c.drawString(340, h - 120, "RIGHT two")
    c.save()
    return p


@pytest.fixture
def cross_page_table_pdf(tmp_path):
    """2페이지에 걸친 연속 표 PDF (개발계획서 §6.1 페이지 횡단 병합)."""
    p = tmp_path / "cross_page.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    w, h = A4
    # 1페이지: 헤더 + 데이터 일부 (하단)
    t1 = Table([["Item", "Val"], ["A", "1"], ["B", "2"]],
               colWidths=[100, 60], rowHeights=22)
    t1.setStyle(_grid_style())
    t1.wrapOn(c, w, h)
    t1.drawOn(c, 72, 120)
    c.showPage()
    # 2페이지: 같은 열 수의 데이터 연속 (상단)
    t2 = Table([["C", "3"], ["D", "4"]], colWidths=[100, 60], rowHeights=22)
    t2.setStyle(_grid_style())
    t2.wrapOn(c, w, h)
    t2.drawOn(c, 72, h - 160)
    c.save()
    return p


def _grid_style():
    from reportlab.platypus import TableStyle
    from reportlab.lib import colors
    return TableStyle([
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("FONTSIZE", (0, 0), (-1, -1), 11),
    ])


@pytest.fixture
def docx_file(tmp_path):
    """제목·강조·목록·표가 섞인 DOCX (FR-2)."""
    import docx
    p = tmp_path / "sample.docx"
    d = docx.Document()
    d.add_heading("문서 제목", level=1)
    d.add_heading("개요", level=2)
    para = d.add_paragraph("이 문장은 ")
    para.add_run("굵게").bold = True
    para.add_run(" 강조됩니다.")
    d.add_paragraph("첫째 항목", style="List Bullet")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "헤더A"; t.cell(0, 1).text = "헤더B"
    t.cell(1, 0).text = "값1"; t.cell(1, 1).text = "값2"
    d.save(str(p))
    return p


@pytest.fixture
def xlsx_file(tmp_path):
    """병합셀 + 다중시트 XLSX (FR-3)."""
    import openpyxl
    p = tmp_path / "sample.xlsx"
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "매출"
    ws1["A1"] = "구분"; ws1["B1"] = "2024"; ws1["C1"] = "2025"
    ws1["A2"] = "국내"; ws1["B2"] = 120; ws1["C2"] = 150
    ws1.merge_cells("A4:C4")
    ws1["A4"] = "병합제목"
    ws2 = wb.create_sheet("비용")
    ws2["A1"] = "항목"; ws2["B1"] = "금액"
    ws2["A2"] = "임대"; ws2["B2"] = 50
    wb.save(str(p))
    return p
